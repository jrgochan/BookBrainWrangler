"""
Test script for DOCX extraction.
This will create a test DOCX file and try to extract text from it.
"""

import os
import sys
import logging

# Configure logging
logging.basicConfig(level=logging.DEBUG, 
                   format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def create_test_docx(output_path):
    """Create a test DOCX file with known content."""
    try:
        import docx
        from docx.shared import Inches
        
        # Create a new Document
        doc = docx.Document()
        
        # Add a title
        doc.add_heading('Test Document', 0)
        
        # Add some paragraphs
        doc.add_paragraph('This is a test document to verify DOCX extraction.')
        doc.add_paragraph('It contains multiple paragraphs with different styles.')
        
        # Add a heading
        doc.add_heading('Section 1', level=1)
        
        # Add more paragraphs
        p = doc.add_paragraph('This paragraph has ')
        p.add_run('bold').bold = True
        p.add_run(' and ')
        p.add_run('italic').italic = True
        p.add_run(' text.')
        
        # Add a table
        table = doc.add_table(rows=3, cols=3)
        
        # Add content to the table
        table.cell(0, 0).text = 'Row 1, Col 1'
        table.cell(0, 1).text = 'Row 1, Col 2'
        table.cell(0, 2).text = 'Row 1, Col 3'
        table.cell(1, 0).text = 'Row 2, Col 1'
        table.cell(1, 1).text = 'Row 2, Col 2'
        table.cell(1, 2).text = 'Row 2, Col 3'
        table.cell(2, 0).text = 'Row 3, Col 1'
        table.cell(2, 1).text = 'Row 3, Col 2'
        table.cell(2, 2).text = 'Row 3, Col 3'
        
        # Add another heading
        doc.add_heading('Section 2', level=1)
        
        # Add a list
        doc.add_paragraph('Here is a list:', style='List Bullet')
        doc.add_paragraph('Item 1', style='List Bullet')
        doc.add_paragraph('Item 2', style='List Bullet')
        doc.add_paragraph('Item 3', style='List Bullet')
        
        # Save the document
        doc.save(output_path)
        logger.info(f"Created test DOCX file at {output_path}")
        return True
    except Exception as e:
        logger.error(f"Error creating test DOCX: {str(e)}")
        return False

def test_docx_extraction(file_path):
    """Test extracting text from a DOCX file."""
    try:
        from document_processing.formats.docx import DOCXProcessor
        
        # Create a processor
        processor = DOCXProcessor()
        
        # Check if processor is available
        if not processor.is_available():
            logger.error("DOCX processor is not available - python-docx might be missing")
            return False
        
        # Process the file
        logger.info(f"Processing DOCX file: {file_path}")
        result = processor.process(file_path)
        
        # Check the result
        logger.info(f"Extracted {len(result.get('text', ''))} characters of text")
        logger.info(f"Extracted {len(result.get('images', []))} images")
        
        # Print a sample of the text
        text = result.get('text', '')
        print("\n--- First 500 characters of extracted text ---")
        print(text[:500])
        print("----------------------------------------------\n")
        
        return len(text) > 0
    except Exception as e:
        logger.error(f"Error in test_docx_extraction: {str(e)}")
        import traceback
        logger.error(f"Traceback: {traceback.format_exc()}")
        return False

if __name__ == "__main__":
    test_file = "test_document.docx"
    
    # Create the test file
    if create_test_docx(test_file):
        # Test extraction
        if test_docx_extraction(test_file):
            logger.info("DOCX extraction test passed!")
        else:
            logger.error("DOCX extraction test failed!")
    else:
        logger.error("Failed to create test DOCX file")
