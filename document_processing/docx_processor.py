"""
DOCX processing utilities for document extraction.
"""

import os
import io
import base64
import pathlib
import requests
from typing import Dict, List, Any, Optional, Callable
import docx
from PIL import Image

from utils.logger import get_logger

# Get a logger for this module
logger = get_logger(__name__)

def process_docx(docx_path: str, include_images: bool = True, progress_callback: Optional[Callable] = None) -> Dict[str, Any]:
    """
    Process a DOCX file and extract text and optionally images.
    
    Args:
        docx_path: Path to the DOCX file
        include_images: Whether to include images in the extraction
        progress_callback: Optional callback function for progress updates
        
    Returns:
        A dictionary with extracted content:
        {
            'text': Extracted text as a string,
            'images': List of image descriptions with embedded base64 data (if include_images=True)
        }
    """
    # Initialize result
    result = {
        'text': '',
        'images': []
    }
    
    # Define a helper function for sending progress updates
    def send_progress(current, total, message):
        if progress_callback:
            progress_callback(current, total, message)
    
    try:
        # Load the document
        document = docx.Document(docx_path)
        
        # Count total paragraphs and tables for progress reporting
        total_elements = len(document.paragraphs) + len(document.tables)
        send_progress(0, total_elements, "Starting DOCX processing...")
        
        # Extract text from paragraphs
        for i, paragraph in enumerate(document.paragraphs):
            para_text = paragraph.text
            result['text'] += para_text + "\n"
            
            # Report progress
            progress_text = f"Processing paragraph {i+1}/{len(document.paragraphs)}"
            send_progress(i, total_elements, {'text': progress_text, 'action': 'processing'})
        
        # Extract text from tables
        for i, table in enumerate(document.tables):
            table_text = ""
            for row in table.rows:
                row_text = [cell.text for cell in row.cells]
                table_text += " | ".join(row_text) + "\n"
            
            result['text'] += table_text + "\n\n"
            
            # Report progress
            progress_text = f"Processing table {i+1}/{len(document.tables)}"
            send_progress(len(document.paragraphs) + i, total_elements, 
                         {'text': progress_text, 'action': 'processing'})
        
        # Extract images if requested
        if include_images:
            extract_images_from_docx(docx_path, document, result, progress_callback)
        
        return result
        
    except Exception as e:
        logger.error(f"Error processing DOCX file: {str(e)}")
        return result

def extract_images_from_docx(docx_path: str, document: Any, result: Dict[str, Any], 
                            progress_callback: Optional[Callable] = None) -> None:
    """
    Extract images from a DOCX document and add them to the result.
    
    Args:
        docx_path: Path to the DOCX file
        document: Loaded docx.Document object
        result: Result dictionary to add images to
        progress_callback: Optional callback function for progress updates
    """
    try:
        # Get the relationship IDs for images
        image_rels = []
        for rel in document.part.rels.values():
            if "image" in rel.target_ref:
                image_rels.append(rel)
        
        # Process each image
        for i, rel in enumerate(image_rels):
            try:
                # Check if this is an external image
                is_external = hasattr(rel, 'target_mode') and rel.target_mode == 'External'
                
                image_data = None
                if is_external:
                    # Handle external image reference
                    target_url = rel.target_ref
                    logger.info(f"Processing external image {i+1}: {target_url}")
                    
                    # Handle different external reference types
                    if target_url.startswith(('http://', 'https://')):
                        # Download from URL
                        try:
                            response = requests.get(target_url, timeout=10)
                            response.raise_for_status()  # Raise exception for HTTP errors
                            image_data = response.content
                            logger.info(f"Successfully downloaded image from URL: {target_url}")
                        except Exception as e:
                            logger.error(f"Failed to download image from URL {target_url}: {str(e)}")
                            continue
                    else:
                        # Handle file path (could be absolute or relative)
                        try:
                            # First try as absolute path
                            img_path = pathlib.Path(target_url)
                            
                            # If not exists and seems like a relative path, try relative to the docx
                            if not img_path.exists() and not img_path.is_absolute():
                                docx_dir = pathlib.Path(docx_path).parent
                                img_path = docx_dir / target_url
                            
                            # If we found the file, read it
                            if img_path.exists():
                                with open(img_path, 'rb') as f:
                                    image_data = f.read()
                                logger.info(f"Successfully read image from file: {img_path}")
                            else:
                                logger.error(f"External image file not found: {target_url}")
                                continue
                        except Exception as e:
                            logger.error(f"Error loading external image file: {str(e)}")
                            continue
                else:
                    # Internal image
                    try:
                        image_part = rel.target_part
                        image_data = image_part.blob
                        logger.info(f"Successfully extracted internal image {i+1}")
                    except Exception as e:
                        logger.error(f"Error extracting internal image {i+1}: {str(e)}")
                        continue
                
                # Process the image data if we have it
                if image_data:
                    try:
                        # Convert to PIL Image
                        from io import BytesIO
                        img = Image.open(BytesIO(image_data))
                        
                        # Convert to base64
                        img_byte_arr = BytesIO()
                        img.save(img_byte_arr, format='PNG')
                        img_base64 = base64.b64encode(img_byte_arr.getvalue()).decode('utf-8')
                        
                        # Add to result
                        result['images'].append({
                            'page': 1,  # DOCX doesn't have pages like PDF
                            'index': i,
                            'description': f"Image {i+1} from document",
                            'data': img_base64
                        })
                        
                        logger.info(f"Successfully added image {i+1} to result")
                    except Exception as e:
                        logger.error(f"Error processing image data: {str(e)}")
            except Exception as e:
                logger.error(f"Error processing image {i+1}: {str(e)}")
                continue
                
    except Exception as e:
        logger.error(f"Error extracting images from DOCX: {str(e)}")