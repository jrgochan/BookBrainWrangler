"""
DOCX processing module for document extraction.

This module provides functionality for extracting text and images from DOCX files.
"""

import os
import io
import base64
import tempfile
from typing import Dict, List, Any, Optional, Union, Callable, Tuple

from utils.logger import get_logger
from core.exceptions import DocumentProcessingError, DocumentFormatError

# Initialize logger
logger = get_logger(__name__)

# Try to import required libraries
try:
    import docx
    from docx.document import Document
    from docx.oxml.text.paragraph import CT_P
    from docx.text.paragraph import Paragraph
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    DOCX_AVAILABLE = True
    logger.info("Successfully imported python-docx and related modules")
except ImportError as e:
    logger.warning(f"python-docx not available. DOCX processing will be limited. Error: {str(e)}")
    DOCX_AVAILABLE = False
except Exception as e:
    logger.error(f"Error importing python-docx: {str(e)}")
    DOCX_AVAILABLE = False

class DOCXProcessor:
    """
    DOCX processor for extracting text and images from DOCX files.
    """
    
    def __init__(self):
        """Initialize the DOCX processor."""
        # Check dependencies
        if not DOCX_AVAILABLE:
            logger.warning("python-docx not available. DOCX processing will be limited.")
            
    def is_available(self) -> bool:
        """
        Check if the DOCX processor is available.
        
        Returns:
            True if the processor is fully available, False otherwise
        """
        return DOCX_AVAILABLE
    
    def process(
        self, 
        file_path: str, 
        extract_images: bool = True,
        progress_callback: Optional[Callable] = None
    ) -> Dict[str, Any]:
        """
        Process a DOCX file and extract text and optionally images.
        
        Args:
            file_path: Path to the DOCX file
            extract_images: Whether to extract images
            progress_callback: Optional callback function for progress updates
            
        Returns:
            A dictionary with extracted content:
            {
                'text': Extracted text as a string,
                'images': List of image descriptions with embedded base64 data (if extract_images=True)
            }
        """
        logger.info(f"Starting DOCX processing for file: {file_path}")
        
        if not os.path.exists(file_path):
            error_msg = f"DOCX file not found: {file_path}"
            logger.error(error_msg)
            raise FileNotFoundError(error_msg)
        
        if not file_path.lower().endswith(('.docx', '.doc')):
            error_msg = f"Not a DOCX file: {file_path}"
            logger.error(error_msg)
            raise DocumentFormatError(error_msg)
        
        # Initialize result dictionary
        result = {
            'text': '',
            'images': []
        }
        
        try:
            # Define a progress callback handler
            def send_progress(current, total, message="Processing DOCX"):
                if progress_callback:
                    if callable(progress_callback):
                        progress_callback(current / total, message)
                    else:
                        logger.warning(f"Progress callback is not callable: {type(progress_callback)}")
                    
            # Extract text from DOCX
            if DOCX_AVAILABLE:
                logger.info(f"python-docx is available, processing file: {file_path}")
                # Load the document
                try:
                    doc = docx.Document(file_path)
                    logger.info(f"Document loaded successfully, found {len(doc.paragraphs)} paragraphs and {len(doc.tables)} tables")
                except Exception as e:
                    logger.error(f"Failed to load document: {str(e)}")
                    raise
                
                # Extract text
                text_result = self.extract_text(doc, progress_callback=send_progress)
                result['text'] = text_result.get('text', '')
                logger.info(f"Text extraction completed, extracted {len(result['text'])} characters")
                
                # Extract images if requested
                if extract_images:
                    logger.info("Starting image extraction")
                    images_result = self.extract_images(file_path, doc, progress_callback=send_progress)
                    result['images'] = images_result.get('images', [])
                    logger.info(f"Image extraction completed, extracted {len(result['images'])} images")
            else:
                logger.warning("python-docx not available, skipping DOCX processing")
                                
            # Log success
            logger.info(f"Successfully processed DOCX file: {file_path}")
            logger.info(f"Result summary: {len(result['text'])} characters of text, {len(result['images'])} images")
            return result
            
        except Exception as e:
            error_msg = f"Error processing DOCX file {file_path}: {str(e)}"
            logger.error(error_msg)
            raise DocumentProcessingError(error_msg) from e
    
    def extract_text(
        self, 
        document,
        progress_callback: Optional[Callable] = None
    ) -> Dict[str, Any]:
        """
        Extract text from a DOCX document.
        
        Args:
            document: Loaded docx.Document object
            progress_callback: Optional callback function for progress updates
            
        Returns:
            Dictionary with extracted text
        """
        if not DOCX_AVAILABLE:
            logger.error("DOCX processing not available - python-docx module missing")
            return {'text': ''}
        
        try:
            # Get paragraphs count for progress tracking
            paragraphs = document.paragraphs
            total_paragraphs = len(paragraphs)
            
            logger.info(f"Starting text extraction from DOCX with {total_paragraphs} paragraphs and {len(document.tables)} tables")
            
            # Extract text from paragraphs
            text_parts = []
            for i, paragraph in enumerate(paragraphs):
                # Update progress
                if progress_callback:
                    try:
                        # Ensure i and total_paragraphs are properly converted to float
                        i_float = float(i)
                        # Ensure total_paragraphs is a number and not a string
                        if isinstance(total_paragraphs, str):
                            logger.debug(f"total_paragraphs is a string: '{total_paragraphs}', converting to int")
                            total_para_num = float(int(total_paragraphs))
                        else:
                            total_para_num = float(total_paragraphs)
                            
                        # Calculate progress - use 1.0 if we can't determine valid progress
                        if total_para_num > 0:
                            progress_fraction = i_float / total_para_num
                        else:
                            progress_fraction = 0.0
                            
                        progress_callback(progress_fraction, f"Extracting text from paragraph {i+1}/{total_paragraphs}")
                    except Exception as e:
                        # Don't let progress reporting stop the extraction
                        logger.debug(f"Progress callback error: {str(e)}")
                        logger.debug(f"Progress callback error details - i: {type(i)}={i}, total_paragraphs: {type(total_paragraphs)}={total_paragraphs}")
                
                # Add paragraph text
                para_text = paragraph.text
                logger.debug(f"Paragraph {i+1} text: {para_text[:50]}{'...' if len(para_text) > 50 else ''}")
                if para_text.strip():  # Only add non-empty paragraphs
                    text_parts.append(para_text)
            
            logger.info(f"Extracted {len(text_parts)} non-empty paragraphs")
            
            # Get tables content
            table_parts = []
            for i, table in enumerate(document.tables):
                logger.debug(f"Processing table {i+1} with {len(table.rows)} rows")
                for row_idx, row in enumerate(table.rows):
                    row_text = []
                    for cell in row.cells:
                        # For each cell, also process paragraphs within it
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    
                    if row_text:  # Only add non-empty rows
                        table_text = ' | '.join(row_text)
                        logger.debug(f"Table {i+1}, Row {row_idx+1}: {table_text[:50]}{'...' if len(table_text) > 50 else ''}")
                        table_parts.append(table_text)
            
            logger.info(f"Extracted {len(table_parts)} non-empty table rows")
            
            # Add table parts to text parts
            text_parts.extend(table_parts)
            
            # Join all text parts with double newlines
            full_text = '\n\n'.join(text_parts)
            
            # Log success with character count
            char_count = len(full_text)
            logger.info(f"Extracted text from DOCX: {char_count} characters from {len(text_parts)} text parts")
            
            if char_count == 0:
                logger.warning("DOCX text extraction produced 0 characters - this may indicate an issue with the document format or content")
            
            return {'text': full_text.strip()}
            
        except Exception as e:
            import traceback
            logger.error(f"Error extracting text from DOCX: {str(e)}")
            logger.error(f"Error traceback: {traceback.format_exc()}")
            # Return empty text but with error indicator
            return {'text': '', 'error': str(e)}
    
    def extract_images(
        self,
        file_path: str,
        document,
        progress_callback: Optional[Callable] = None
    ) -> Dict[str, Any]:
        """
        Extract images from a DOCX file.
        
        Args:
            file_path: Path to the DOCX file
            document: Loaded docx.Document object
            progress_callback: Optional callback function for progress updates
            
        Returns:
            Dictionary with list of images as base64 encoded strings
        """
        if not DOCX_AVAILABLE:
            return {'images': []}
        
        try:
            # Get document relationships
            doc_part = document.part
            image_parts = []
            
            # Find image relationships
            for rel in doc_part.rels.values():
                if rel.reltype == RT.IMAGE:
                    try:
                        # Check if this is an internal or external image
                        if hasattr(rel, 'target_mode') and rel.target_mode == 'External':
                            logger.debug(f"Found external image reference: {rel.target_ref}")
                            # Skip external images for now as they require special handling
                            continue
                        else:
                            # This is an internal image
                            image_parts.append(rel.target_part)
                    except Exception as e:
                        logger.debug(f"Skipping problematic image relationship: {str(e)}")
                        continue
            
            # Extract images
            images = []
            total_images = len(image_parts)
            
            for i, image_part in enumerate(image_parts):
                # Update progress
                if progress_callback and total_images > 0:
                    try:
                        i_float = float(i)
                        total_float = float(total_images)
                        progress_callback(i_float / total_float, f"Extracting image {i+1}/{total_images}")
                    except Exception as e:
                        logger.debug(f"Progress callback error in image extraction: {str(e)}")
                
                # Get image data
                image_data = image_part.blob
                
                # Get image format
                image_format = image_part.content_type.split('/')[-1]
                if image_format == 'jpeg':
                    image_format = 'jpg'
                
                # Encode image data as base64
                img_base64 = base64.b64encode(image_data).decode('utf-8')
                
                # Add to result
                images.append({
                    'index': i,
                    'type': 'embedded',
                    'format': image_format,
                    'data': img_base64
                })
            
            # Log success
            logger.info(f"Extracted {len(images)} images from DOCX file")
            return {'images': images}
            
        except Exception as e:
            logger.error(f"Error extracting images from DOCX: {str(e)}")
            return {'images': []}
    
    def get_thumbnail(
        self,
        file_path: str,
        width: int = 200,
        height: int = 300
    ) -> Optional[str]:
        """
        Get a thumbnail image for a DOCX file.
        Uses the first image in the document if available.
        
        Args:
            file_path: Path to the DOCX file
            width: Desired width of the thumbnail
            height: Desired height of the thumbnail
            
        Returns:
            Base64 encoded thumbnail image or None if failed
        """
        try:
            if not DOCX_AVAILABLE:
                return None
                
            # Process the document
            doc = docx.Document(file_path)
            
            # Extract images
            images_result = self.extract_images(file_path, doc)
            images = images_result.get('images', [])
            
            if not images:
                # No images found, try to create a text thumbnail
                return self._create_text_thumbnail(doc, width, height)
            
            # Use the first image as thumbnail
            image_data = images[0]['data']
            image_format = images[0]['format']
            
            # Create a thumbnail using PIL
            try:
                from PIL import Image
                import io
                import base64
                
                # Decode the base64 image
                image_bytes = base64.b64decode(image_data)
                
                # Open the image
                image = Image.open(io.BytesIO(image_bytes))
                
                # Resize the image while maintaining aspect ratio
                original_width, original_height = image.size
                aspect_ratio = original_width / original_height
                
                if width / height > aspect_ratio:
                    # Height is the limiting factor
                    new_height = height
                    new_width = int(aspect_ratio * new_height)
                else:
                    # Width is the limiting factor
                    new_width = width
                    new_height = int(new_width / aspect_ratio)
                
                # Resize the image
                resized_image = image.resize((new_width, new_height), Image.LANCZOS)
                
                # Create a new image with the desired dimensions and paste the resized image
                thumbnail = Image.new('RGB', (width, height), (255, 255, 255))
                paste_x = (width - new_width) // 2
                paste_y = (height - new_height) // 2
                thumbnail.paste(resized_image, (paste_x, paste_y))
                
                # Save the thumbnail to a byte buffer
                buffered = io.BytesIO()
                thumbnail.save(buffered, format="JPEG", quality=85)
                
                # Get base64 encoded string
                thumbnail_base64 = base64.b64encode(buffered.getvalue()).decode('utf-8')
                
                return thumbnail_base64
                
            except ImportError:
                logger.warning("PIL not available for thumbnail generation")
                return None
                
        except Exception as e:
            logger.error(f"Error creating thumbnail for {file_path}: {str(e)}")
            return None
    
    def _create_text_thumbnail(
        self,
        document,
        width: int = 200,
        height: int = 300
    ) -> Optional[str]:
        """
        Create a text-based thumbnail for a DOCX document.
        
        Args:
            document: Loaded docx.Document object
            width: Desired width of the thumbnail
            height: Desired height of the thumbnail
            
        Returns:
            Base64 encoded thumbnail image or None if failed
        """
        try:
            # Get document title or first paragraph
            title = None
            for para in document.paragraphs:
                if para.text.strip():
                    title = para.text.strip()
                    break
            
            if not title:
                title = "DOCX Document"
                
            # Create text thumbnail using PIL
            try:
                from PIL import Image, ImageDraw, ImageFont
                import io
                import base64
                import tempfile
                
                # Create a new image
                image = Image.new('RGB', (width, height), (245, 245, 245))
                draw = ImageDraw.Draw(image)
                
                # Draw background
                draw.rectangle([(0, 0), (width, 40)], fill=(59, 89, 152))  # Header color
                
                # Draw DOCX icon
                icon_margin = 20
                icon_size = min(width - 2*icon_margin, height - 2*icon_margin - 40)
                icon_box = (
                    (width - icon_size) // 2,
                    ((height - 40) - icon_size) // 2 + 40
                )
                draw.rectangle(
                    [icon_box, (icon_box[0] + icon_size, icon_box[1] + icon_size)],
                    fill=(255, 255, 255),
                    outline=(200, 200, 200)
                )
                
                # Draw document lines
                line_width = int(icon_size * 0.7)
                line_x = (width - line_width) // 2
                line_y_start = icon_box[1] + int(icon_size * 0.2)
                line_height = int(icon_size * 0.1)
                line_gap = int(icon_size * 0.05)
                
                for i in range(4):
                    y = line_y_start + i * (line_height + line_gap)
                    draw.rectangle(
                        [(line_x, y), (line_x + line_width, y + line_height)],
                        fill=(220, 220, 220)
                    )
                
                # Draw title text
                try:
                    # Try to load a font
                    font = ImageFont.truetype("arial.ttf", 12)
                except:
                    # Use default font
                    font = ImageFont.load_default()
                
                # Truncate title if it's too long
                max_chars = 20
                if len(title) > max_chars:
                    title = title[:max_chars-3] + "..."
                
                # Get text size
                text_width = draw.textlength(title, font=font)
                text_x = (width - text_width) // 2
                
                # Draw title
                draw.text((text_x, 15), title, font=font, fill=(255, 255, 255))
                
                # Draw format indicator
                format_text = "DOCX"
                format_width = draw.textlength(format_text, font=font)
                format_x = (width - format_width) // 2
                format_y = height - 25
                draw.text((format_x, format_y), format_text, font=font, fill=(100, 100, 100))
                
                # Save the thumbnail to a byte buffer
                buffered = io.BytesIO()
                image.save(buffered, format="JPEG", quality=85)
                
                # Get base64 encoded string
                thumbnail_base64 = base64.b64encode(buffered.getvalue()).decode('utf-8')
                
                return thumbnail_base64
                
            except ImportError:
                logger.warning("PIL not available for text thumbnail generation")
                return None
                
        except Exception as e:
            logger.error(f"Error creating text thumbnail: {str(e)}")
            return None
