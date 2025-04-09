"""
Metadata extraction module for Book Knowledge AI.
Extracts metadata from various document formats.
"""

import os
import re
from datetime import datetime
from typing import Dict, Any, Optional, List

from utils.logger import get_logger

# Get a logger for this module
logger = get_logger(__name__)

def extract_metadata(file_path: str) -> Dict[str, Any]:
    """
    Extract metadata from a document.
    Dispatches to format-specific extractors based on file extension.
    
    Args:
        file_path: Path to the document
        
    Returns:
        Dictionary of metadata
    """
    if not os.path.exists(file_path):
        logger.error(f"File not found: {file_path}")
        return {}
    
    # Get file extension
    _, ext = os.path.splitext(file_path)
    ext = ext.lower().lstrip(".")
    
    # Basic metadata for all files
    metadata = {
        "filename": os.path.basename(file_path),
        "file_path": file_path,
        "file_size": os.path.getsize(file_path),
        "file_extension": ext,
        "format": ext,
        "last_modified": datetime.fromtimestamp(os.path.getmtime(file_path)).strftime("%Y-%m-%d %H:%M:%S"),
        "extracted_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    }
    
    # Format-specific metadata
    if ext == "pdf":
        pdf_metadata = extract_pdf_metadata(file_path)
        metadata.update(pdf_metadata)
    elif ext == "docx":
        docx_metadata = extract_docx_metadata(file_path)
        metadata.update(docx_metadata)
    elif ext == "txt":
        txt_metadata = extract_txt_metadata(file_path)
        metadata.update(txt_metadata)
    
    return metadata

def extract_pdf_metadata(file_path: str) -> Dict[str, Any]:
    """
    Extract metadata from a PDF document.
    
    Args:
        file_path: Path to the PDF document
        
    Returns:
        Dictionary of metadata
    """
    try:
        import PyPDF2
        
        metadata = {}
        
        with open(file_path, 'rb') as f:
            pdf = PyPDF2.PdfReader(f)
            
            # Basic metadata
            metadata["page_count"] = len(pdf.pages)
            metadata["format"] = "PDF"
            
            # Document info dictionary
            if pdf.metadata:
                info = pdf.metadata
                
                # Extract standard metadata fields
                if info.get("/Title"):
                    metadata["title"] = info.get("/Title")
                if info.get("/Author"):
                    metadata["author"] = info.get("/Author")
                if info.get("/Subject"):
                    metadata["subject"] = info.get("/Subject")
                if info.get("/Keywords"):
                    metadata["keywords"] = info.get("/Keywords")
                if info.get("/Producer"):
                    metadata["producer"] = info.get("/Producer")
                if info.get("/Creator"):
                    metadata["creator"] = info.get("/Creator")
                if info.get("/CreationDate"):
                    creation_date = info.get("/CreationDate")
                    if isinstance(creation_date, str) and creation_date.startswith("D:"):
                        # Convert PDF date format to readable date
                        date_str = creation_date[2:14]  # Extract YYYYMMDDHHMMSS
                        try:
                            date = datetime.strptime(date_str, "%Y%m%d%H%M%S")
                            metadata["creation_date"] = date.strftime("%Y-%m-%d %H:%M:%S")
                        except ValueError:
                            metadata["creation_date"] = creation_date
                    else:
                        metadata["creation_date"] = creation_date
            
            # Try to extract TOC (table of contents)
            try:
                toc = pdf.outline
                if toc:
                    metadata["has_toc"] = True
            except Exception:
                metadata["has_toc"] = False
        
        # Try to extract document title from first page if not found in metadata
        if "title" not in metadata or not metadata["title"]:
            first_page_text = pdf.pages[0].extract_text()
            title = extract_title_from_text(first_page_text)
            if title:
                metadata["title"] = title
        
        return metadata
    
    except Exception as e:
        logger.error(f"Error extracting PDF metadata: {str(e)}")
        return {"format": "PDF", "error": str(e)}

def extract_docx_metadata(file_path: str) -> Dict[str, Any]:
    """
    Extract metadata from a DOCX document.
    
    Args:
        file_path: Path to the DOCX document
        
    Returns:
        Dictionary of metadata
    """
    try:
        import docx
        
        metadata = {}
        
        doc = docx.Document(file_path)
        
        # Basic metadata
        metadata["format"] = "DOCX"
        metadata["paragraph_count"] = len(doc.paragraphs)
        
        # Core properties
        try:
            core_props = doc.core_properties
            
            if core_props.title:
                metadata["title"] = core_props.title
            if core_props.author:
                metadata["author"] = core_props.author
            if core_props.subject:
                metadata["subject"] = core_props.subject
            if core_props.keywords:
                metadata["keywords"] = core_props.keywords
            if core_props.category:
                metadata["category"] = core_props.category
            if core_props.comments:
                metadata["comments"] = core_props.comments
            if core_props.language:
                metadata["language"] = core_props.language
            if core_props.created:
                metadata["creation_date"] = core_props.created.strftime("%Y-%m-%d %H:%M:%S")
            if core_props.modified:
                metadata["modified_date"] = core_props.modified.strftime("%Y-%m-%d %H:%M:%S")
            if core_props.last_modified_by:
                metadata["last_modified_by"] = core_props.last_modified_by
            if core_props.revision:
                metadata["revision"] = core_props.revision
            if core_props.version:
                metadata["version"] = core_props.version
        except Exception as e:
            logger.warning(f"Error extracting DOCX core properties: {str(e)}")
        
        # Document statistics
        try:
            metadata["section_count"] = len(doc.sections)
            metadata["table_count"] = len(doc.tables)
            
            # Count images
            image_count = 0
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    image_count += 1
            metadata["image_count"] = image_count
        except Exception as e:
            logger.warning(f"Error extracting DOCX document statistics: {str(e)}")
        
        # Try to extract document title from first paragraphs if not found in metadata
        if "title" not in metadata or not metadata["title"]:
            first_paragraphs_text = "\n".join([p.text for p in doc.paragraphs[:5]])
            title = extract_title_from_text(first_paragraphs_text)
            if title:
                metadata["title"] = title
        
        return metadata
    
    except Exception as e:
        logger.error(f"Error extracting DOCX metadata: {str(e)}")
        return {"format": "DOCX", "error": str(e)}

def extract_txt_metadata(file_path: str) -> Dict[str, Any]:
    """
    Extract metadata from a text document.
    
    Args:
        file_path: Path to the text document
        
    Returns:
        Dictionary of metadata
    """
    try:
        metadata = {}
        
        # Basic metadata
        metadata["format"] = "TXT"
        
        # Read text content
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()
        
        # Text statistics
        metadata["character_count"] = len(content)
        metadata["line_count"] = content.count('\n') + 1
        metadata["word_count"] = len(content.split())
        
        # Try to extract document title from first lines
        title = extract_title_from_text(content[:1000])
        if title:
            metadata["title"] = title
        
        return metadata
    
    except Exception as e:
        logger.error(f"Error extracting TXT metadata: {str(e)}")
        return {"format": "TXT", "error": str(e)}

def extract_title_from_text(text: str) -> Optional[str]:
    """
    Extract a potential title from text content.
    
    Args:
        text: Text content
        
    Returns:
        Potential document title or None
    """
    if not text:
        return None
    
    # Split into lines
    lines = text.strip().split('\n')
    
    # Find first non-empty line
    for line in lines:
        line = line.strip()
        if line and len(line) > 3 and len(line) < 100:
            # Check if it looks like a title (not too long, not too short, etc.)
            return line
    
    return None