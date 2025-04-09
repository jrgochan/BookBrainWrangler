"""
Document processor module for handling document content extraction.
"""

import os
import io
import re
import base64
import json
import time
import requests
import numpy as np
from urllib.parse import urlparse
import pathlib
from typing import Dict, List, Any, Optional, Union, Tuple, Callable
import PyPDF2
import pytesseract
from pdf2image import convert_from_path
import docx
from PIL import Image

# Import EasyOCR conditionally to handle environments where it might not be available
try:
    import easyocr
    EASYOCR_AVAILABLE = True
except ImportError:
    EASYOCR_AVAILABLE = False

from utils.logger import get_logger

# Get a logger for this module
logger = get_logger(__name__)

# OCR Engine options
OCR_ENGINES = {
    "pytesseract": "PyTesseract",
    "easyocr": "EasyOCR"
}

class DocumentProcessor:
    def __init__(self, ocr_engine="pytesseract", ocr_settings=None):
        """
        Initialize the document processor with specified settings.
        
        Args:
            ocr_engine: The OCR engine to use ('pytesseract' or 'easyocr')
            ocr_settings: Dictionary of OCR-specific settings
        """
        # Set default OCR settings if none provided
        self.ocr_settings = ocr_settings or {}
        
        # Set the OCR engine
        self.ocr_engine = ocr_engine if ocr_engine in OCR_ENGINES else "pytesseract"
        
        # Initialize OCR engines conditionally
        if self.ocr_engine == "pytesseract":
            # Auto-detect Tesseract path based on platform
            tesseract_path = self._detect_tesseract_path(self.ocr_settings)
            if tesseract_path:
                pytesseract.pytesseract.tesseract_cmd = tesseract_path
                logger.info(f"Using Tesseract executable at: {tesseract_path}")
            else:
                logger.warning("Tesseract path not found. Using system default, which may not work.")
                
        elif self.ocr_engine == "easyocr" and EASYOCR_AVAILABLE:
            # Initialize EasyOCR reader with languages from settings or default to English
            languages = self.ocr_settings.get('languages', ['en'])
            self.easyocr_reader = easyocr.Reader(languages)
            logger.info(f"Initialized EasyOCR with languages: {languages}")
        
        # Set metadata extraction patterns
        self.metadata_patterns = {
            'title': [
                r'title:\s*([^\n]+)',
                r'^(?:book\s+)?title[:\s]+([^\n]+)',
                r'(?:^|\n)([A-Z][^\n]{5,50})(?:\n|$)',  # Capitalized line that could be a title
            ],
            'author': [
                r'(?:author|by)(?:s)?:\s*([^\n]+)',
                r'(?:written|published)\s+by\s+([^\n]+)',
                r'©\s*([^\n]+)',  # Copyright symbol followed by name
            ],
            'categories': [
                r'(?:category|categories|genre|genres|subject|subjects):\s*([^\n]+)',
                r'(?:keywords|tags):\s*([^\n]+)'
            ]
        }
    
    def extract_content(self, file_path, include_images=True, progress_callback=None):
        """
        Extract text and optionally images from a document file (PDF or DOCX).
        
        Args:
            file_path: Path to the document file
            include_images: Whether to include images in the extraction
            progress_callback: Optional callback function for progress updates
            
        Returns:
            A dictionary with extracted content:
            {
                'text': Extracted text as a string,
                'images': List of image descriptions with embedded base64 data (if include_images=True)
            }
        """
        # Determine file type from extension
        file_ext = os.path.splitext(file_path)[1].lower()
        
        # Extract content based on file type
        if file_ext == '.pdf':
            return self._process_pdf(file_path, include_images, progress_callback)
        elif file_ext == '.docx':
            return self._process_docx(file_path, include_images, progress_callback)
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")
    
    def extract_text(self, file_path, progress_callback=None):
        """
        Legacy method to extract only text from a document file.
        Maintained for backward compatibility.
        
        Args:
            file_path: Path to the document file
            progress_callback: Optional callback function for progress updates
            
        Returns:
            Extracted text as a string
        """
        result = self.extract_content(file_path, include_images=False, progress_callback=progress_callback)
        if isinstance(result, dict):
            return result.get('text', '')
        return result
    
    def _process_pdf(self, pdf_path, include_images=True, progress_callback=None):
        """
        Process a PDF file and extract text and optionally images.
        
        Args:
            pdf_path: Path to the PDF file
            include_images: Whether to include images in the extraction
            progress_callback: Optional callback function for progress updates
            
        Returns:
            A dictionary with extracted content:
            {
                'text': Extracted text as a string,
                'images': List of image descriptions with embedded base64 data (if include_images=True)
            }
        """
        # We're using global imports now
        
        # Define a helper function for sending progress updates
        def send_progress(current, total, message):
            if progress_callback:
                progress_callback(current, total, message)
        
        # Get total pages
        total_pages = self.get_page_count(pdf_path)
        if total_pages == 0:
            return {'text': '', 'images': []}
        
        # Initialize result
        result = {
            'text': '',
            'images': []
        }
        
        send_progress(0, total_pages, "Starting document processing...")
        
        # Open the PDF file
        with open(pdf_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            # Process each page
            for i in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[i]
                progress_text = f"Processing page {i+1}/{total_pages}"
                send_progress(i, total_pages, {'text': progress_text, 'action': 'processing'})
                
                # Try to extract text directly from PDF
                extracted_text = page.extract_text()
                
                # If direct extraction didn't work well, use OCR
                if not extracted_text or len(extracted_text.strip()) < 50:  # Heuristic to determine if text extraction failed
                    try:
                        # Convert PDF page to image
                        images = convert_from_path(pdf_path, first_page=i+1, last_page=i+1)
                        if images:
                            # Use selected OCR engine to extract text from the image
                            img = images[0]
                            
                            # Perform OCR based on selected engine
                            extracted_text, confidence = self._perform_ocr(img)
                            
                            # Save image for UI feedback
                            if include_images or progress_callback:
                                # Convert image to base64 for display
                                img_byte_arr = io.BytesIO()
                                img.save(img_byte_arr, format='PNG')
                                img_base64 = base64.b64encode(img_byte_arr.getvalue()).decode('utf-8')
                                
                                if include_images:
                                    # Add image to result
                                    result['images'].append({
                                        'page': i+1,
                                        'index': 0,  # First image on page
                                        'description': f"Image from page {i+1}",
                                        'data': img_base64
                                    })
                                
                                # Send progress with image and text data
                                if progress_callback:
                                    send_progress(i, total_pages, {
                                        'text': progress_text,
                                        'current_image': img_base64,
                                        'ocr_text': extracted_text,
                                        'confidence': confidence,
                                        'action': 'extracting'
                                    })
                    except Exception as e:
                        logger.error(f"OCR processing error on page {i+1}: {str(e)}")
                        # If OCR fails, use whatever text we could extract directly
                        if progress_callback:
                            send_progress(i, total_pages, {
                                'text': f"Error processing page {i+1}: {str(e)}",
                                'action': 'error'
                            })
                else:
                    # Direct extraction worked, report good confidence
                    confidence = 95.0  # High confidence for direct extraction
                    
                    # If we need to show progress with the page image
                    if progress_callback:
                        try:
                            # Generate image from page for visualization
                            images = convert_from_path(pdf_path, first_page=i+1, last_page=i+1)
                            if images:
                                img = images[0]
                                img_byte_arr = io.BytesIO()
                                img.save(img_byte_arr, format='PNG')
                                img_base64 = base64.b64encode(img_byte_arr.getvalue()).decode('utf-8')
                                
                                send_progress(i, total_pages, {
                                    'text': progress_text,
                                    'current_image': img_base64,
                                    'ocr_text': extracted_text,
                                    'confidence': confidence,
                                    'action': 'extracting'
                                })
                        except Exception as e:
                            # If image conversion fails, just report the text
                            logger.error(f"Image conversion error on page {i+1}: {str(e)}")
                            send_progress(i, total_pages, {
                                'text': progress_text,
                                'ocr_text': extracted_text,
                                'confidence': confidence,
                                'action': 'extracting'
                            })
                
                # Add extracted text to result
                result['text'] += extracted_text + "\n\n"
        
        # Finalizing
        send_progress(total_pages-1, total_pages, {'text': "Finalizing document processing...", 'action': 'completed'})
        
        return result
    
    def _process_docx(self, docx_path, include_images=True, progress_callback=None):
        """
        Process a DOCX file and extract text and optionally images.
        
        Args:
            docx_path: Path to the DOCX file
            include_images: Whether to include images in the extraction
            progress_callback: Optional callback function for progress updates
            
        Returns:
            A dictionary with extracted content:
            {
                'text': Extracted text as a string,
                'images': List of image descriptions with embedded base64 data (if include_images=True)
            }
        """
        # We're using global imports now
        
        # Initialize result
        result = {
            'text': '',
            'images': []
        }
        
        # Define a helper function for sending progress updates
        def send_progress(current, total, message):
            if progress_callback:
                progress_callback(current, total, message)
        
        try:
            # Load the document
            document = docx.Document(docx_path)
            
            # Count total paragraphs and tables for progress reporting
            total_elements = len(document.paragraphs) + len(document.tables)
            send_progress(0, total_elements, "Starting DOCX processing...")
            
            # Extract text from paragraphs
            for i, paragraph in enumerate(document.paragraphs):
                para_text = paragraph.text
                result['text'] += para_text + "\n"
                
                # Report progress
                progress_text = f"Processing paragraph {i+1}/{len(document.paragraphs)}"
                send_progress(i, total_elements, {'text': progress_text, 'action': 'processing'})
            
            # Extract text from tables
            for i, table in enumerate(document.tables):
                table_text = ""
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    table_text += " | ".join(row_text) + "\n"
                
                result['text'] += table_text + "\n\n"
                
                # Report progress
                progress_text = f"Processing table {i+1}/{len(document.tables)}"
                send_progress(len(document.paragraphs) + i, total_elements, 
                             {'text': progress_text, 'action': 'processing'})
            
            # Extract images if requested
            if include_images:
                # Get the relationship IDs for images
                image_rels = []
                for rel in document.part.rels.values():
                    if "image" in rel.target_ref:
                        image_rels.append(rel)
                
            # Process each image
            for i, rel in enumerate(image_rels):
                try:
                    # Check if this is an external image
                    is_external = hasattr(rel, 'target_mode') and rel.target_mode == 'External'
                    
                    image_data = None
                    if is_external:
                        # Handle external image reference
                        target_url = rel.target_ref
                        logger.info(f"Processing external image {i+1}: {target_url}")
                        
                        # Handle different external reference types
                        if target_url.startswith(('http://', 'https://')):
                            # Download from URL
                            try:
                                response = requests.get(target_url, timeout=10)
                                response.raise_for_status()  # Raise exception for HTTP errors
                                image_data = response.content
                                logger.info(f"Successfully downloaded image from URL: {target_url}")
                            except Exception as e:
                                logger.error(f"Failed to download image from URL {target_url}: {str(e)}")
                                continue
                        else:
                            # Handle file path (could be absolute or relative)
                            try:
                                # First try as absolute path
                                img_path = pathlib.Path(target_url)
                                
                                # If not exists and seems like a relative path, try relative to the docx
                                if not img_path.exists() and not img_path.is_absolute():
                                    docx_dir = pathlib.Path(docx_path).parent
                                    img_path = docx_dir / target_url
                                
                                # If we found the file, read it
                                if img_path.exists():
                                    with open(img_path, 'rb') as f:
                                        image_data = f.read()
                                    logger.info(f"Successfully read image from file: {img_path}")
                                else:
                                    logger.error(f"External image file not found: {target_url}")
                                    continue
                            except Exception as e:
                                logger.error(f"Failed to read external image file {target_url}: {str(e)}")
                                continue
                    else:
                        # Regular embedded image
                        try:
                            image_data = rel.target_part.blob
                        except Exception as e:
                            logger.error(f"Error accessing embedded image data: {str(e)}")
                            continue
                    
                    # Process image data (same for both external and embedded)
                    if image_data:
                        # Convert to base64 for storage/display
                        img_base64 = base64.b64encode(image_data).decode('utf-8')
                        
                        # Add to result
                        result['images'].append({
                            'page': 1,  # DOCX doesn't have pages like PDF
                            'index': i,
                            'description': f"Image {i+1} from document",
                            'data': img_base64
                        })
                        
                        # For UI display, create an image
                        if progress_callback:
                            try:
                                # Load image for UI display
                                img = Image.open(io.BytesIO(image_data))
                                
                                # Report progress with image
                                progress_text = f"Processing image {i+1}/{len(image_rels)}"
                                send_progress(total_elements - len(image_rels) + i, total_elements, {
                                    'text': progress_text,
                                    'current_image': img_base64,
                                    'ocr_text': f"Image {i+1} from document",
                                    'confidence': 100.0,  # Full confidence for images
                                    'action': 'extracting'
                                })
                            except Exception as e:
                                logger.error(f"Error displaying image {i+1}: {str(e)}")
                except Exception as e:
                    logger.error(f"Error processing image {i+1}: {str(e)}")
            
            # Final progress update
            send_progress(total_elements, total_elements, {'text': "DOCX processing complete", 'action': 'completed'})
            
        except Exception as e:
            error_msg = f"Error processing DOCX file: {str(e)}"
            logger.error(error_msg)
            if progress_callback:
                send_progress(0, 1, {'text': error_msg, 'action': 'error'})
        
        return result
    
    def extract_metadata(self, file_path, content=None):
        """
        Extract book metadata (title, author, categories) from a document.
        
        Args:
            file_path: Path to the document file
            content: Optional pre-extracted content to use instead of re-extracting
            
        Returns:
            A dictionary with metadata fields:
            {
                'title': Extracted title or None if not found,
                'author': Extracted author or None if not found,
                'categories': List of extracted categories or empty list if none found
            }
        """
        # Initialize result
        metadata = {
            'title': None,
            'author': None,
            'categories': []
        }
        
        # Get content if not provided
        if content is None:
            # Only extract the first few pages for metadata (faster)
            try:
                extracted = self.extract_content(file_path, include_images=False)
                if isinstance(extracted, dict):
                    content = extracted.get('text', '')
                else:
                    content = extracted
            except Exception as e:
                logger.error(f"Error extracting content for metadata: {e}")
                return metadata
        
        # Process content to extract metadata
        if content and isinstance(content, str):
            # Extract file name as fallback title (without extension)
            file_name = os.path.basename(file_path)
            file_name_no_ext = os.path.splitext(file_name)[0]
            metadata['title'] = file_name_no_ext  # Default to file name
            
            # Get first 1000 characters for metadata extraction (usually in the beginning)
            # Also check the entire content but prioritize matches near the beginning
            beginning = content[:1000]
            
            # Look for title
            for pattern in self.metadata_patterns['title']:
                # First check beginning of document
                match = re.search(pattern, beginning, re.IGNORECASE)
                if match:
                    metadata['title'] = match.group(1).strip()
                    break
                    
                # If not found in beginning, check full document
                match = re.search(pattern, content, re.IGNORECASE)
                if match:
                    metadata['title'] = match.group(1).strip()
                    break
            
            # Look for author
            for pattern in self.metadata_patterns['author']:
                # Check beginning first
                match = re.search(pattern, beginning, re.IGNORECASE)
                if match:
                    metadata['author'] = match.group(1).strip()
                    break
                    
                # If not found in beginning, check full document
                match = re.search(pattern, content, re.IGNORECASE)
                if match:
                    metadata['author'] = match.group(1).strip()
                    break
            
            # Look for categories
            for pattern in self.metadata_patterns['categories']:
                match = re.search(pattern, content, re.IGNORECASE)
                if match:
                    categories_str = match.group(1).strip()
                    # Split categories by common delimiters
                    categories = re.split(r'[,;|/]', categories_str)
                    metadata['categories'] = [cat.strip() for cat in categories if cat.strip()]
                    break
        
        return metadata
    
    def _detect_tesseract_path(self, settings):
        """
        Auto-detect Tesseract executable path based on platform.
        
        Args:
            settings: Dictionary containing OCR settings
            
        Returns:
            String path to Tesseract executable or None if not found
        """
        # First check if user provided a path in settings
        if settings and 'tesseract_path' in settings:
            user_path = settings['tesseract_path']
            if os.path.exists(user_path):
                logger.info(f"Using user-provided Tesseract path: {user_path}")
                return user_path
        
        # Detect operating system
        import platform
        system = platform.system().lower()
        
        # Define common Tesseract installation paths by platform
        common_paths = {
            'windows': [
                r'C:\Program Files\Tesseract-OCR\tesseract.exe',
                r'C:\Program Files (x86)\Tesseract-OCR\tesseract.exe',
                r'C:\Tesseract-OCR\tesseract.exe',
            ],
            'linux': [
                '/usr/bin/tesseract',
                '/usr/local/bin/tesseract',
            ],
            'darwin': [  # macOS
                '/usr/local/bin/tesseract',
                '/opt/homebrew/bin/tesseract',
                '/opt/local/bin/tesseract',
            ]
        }
        
        # Get list of paths to check for current platform
        paths_to_check = common_paths.get(system, [])
        
        # Check if Tesseract is in the system PATH
        try:
            import shutil
            system_path = shutil.which('tesseract')
            if system_path:
                logger.info(f"Found Tesseract in system PATH: {system_path}")
                return system_path
        except Exception as e:
            logger.warning(f"Error checking system PATH for Tesseract: {str(e)}")
        
        # Check common installation locations
        for path in paths_to_check:
            if os.path.exists(path):
                logger.info(f"Found Tesseract at common location: {path}")
                return path
        
        # If we get here, we couldn't find Tesseract
        logger.warning(f"Could not find Tesseract executable on {system}")
        return None
        
    def _perform_ocr(self, image):
        """
        Perform OCR on an image using the selected OCR engine.
        
        Args:
            image: PIL Image object to perform OCR on
            
        Returns:
            Tuple of (extracted_text, confidence)
        """
        extracted_text = ""
        confidence = 0.0
        
        try:
            if self.ocr_engine == "pytesseract":
                # Use PyTesseract for OCR
                extracted_text = pytesseract.image_to_string(image)
                # Get confidence score
                data = pytesseract.image_to_data(image, output_type=pytesseract.Output.DICT)
                confidences = [float(conf) for conf in data['conf'] if conf != '-1']
                confidence = sum(confidences) / len(confidences) if confidences else 0.0
                logger.debug(f"PyTesseract OCR: {len(extracted_text)} chars, confidence: {confidence:.2f}%")
                
            elif self.ocr_engine == "easyocr" and EASYOCR_AVAILABLE:
                # Use EasyOCR for OCR
                if not hasattr(self, 'easyocr_reader'):
                    # Initialize reader if not already done
                    languages = self.ocr_settings.get('languages', ['en'])
                    self.easyocr_reader = easyocr.Reader(languages)
                    logger.info(f"Initialized EasyOCR with languages: {languages}")
                
                # Convert PIL image to numpy array for EasyOCR
                img_array = np.array(image)
                
                # Perform OCR
                results = self.easyocr_reader.readtext(img_array)
                
                # Extract text and confidence from results
                if results:
                    texts = []
                    confidences = []
                    for detection in results:
                        bbox, text, conf = detection
                        texts.append(text)
                        confidences.append(conf)
                    
                    extracted_text = " ".join(texts)
                    confidence = sum(confidences) / len(confidences) * 100 if confidences else 0.0
                    logger.debug(f"EasyOCR: {len(extracted_text)} chars, confidence: {confidence:.2f}%")
                else:
                    logger.warning("EasyOCR returned no results")
            else:
                # Fallback to PyTesseract if EasyOCR is unavailable or unknown engine specified
                extracted_text = pytesseract.image_to_string(image)
                data = pytesseract.image_to_data(image, output_type=pytesseract.Output.DICT)
                confidences = [float(conf) for conf in data['conf'] if conf != '-1']
                confidence = sum(confidences) / len(confidences) if confidences else 0.0
                logger.warning(f"Fallback to PyTesseract due to unavailability of {self.ocr_engine}")
            
        except Exception as e:
            logger.error(f"Error in OCR processing: {str(e)}")
            # Return empty text and zero confidence on error
            return "", 0.0
            
        return extracted_text, confidence

    def get_page_count(self, pdf_path):
        """
        Get the number of pages in a PDF file.
        
        Args:
            pdf_path: Path to the PDF file
            
        Returns:
            Number of pages as an integer
        """
        # Use PyPDF2 to get the real page count
        try:
            # We're using global imports now
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                return len(pdf_reader.pages)
        except Exception as e:
            logger.error(f"Error getting page count: {str(e)}")
            return 0
